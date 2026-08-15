import argparse
import json
import os
import sys
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TITLE_COLOR = RGBColor(0x8B, 0x1A, 0x1A)
SUBTITLE_COLOR = RGBColor(0x66, 0x66, 0x66)
CAPTION_COLOR = RGBColor(0x88, 0x88, 0x88)
BODY_FONT_SIZE = Pt(12)
LINE_SPACING = 1.5


def set_font(run, ascii_font, east_asia, size=None, bold=None):
    run.font.name = ascii_font
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold


def ensure_docx_supported(path):
    """Convert an image to JPEG if python-docx cannot handle its format (e.g. WEBP)."""
    from PIL import Image

    try:
        with Image.open(path) as im:
            if im.format in ("JPEG", "PNG", "GIF", "BMP", "TIFF"):
                return path
            tmp = os.path.join(tempfile.gettempdir(), "docx_img_" + str(abs(hash(path))) + ".jpg")
            im.convert("RGB").save(tmp, "JPEG", quality=90)
            print(f"converted {path} -> {tmp}")
            return tmp
    except Exception as e:
        print(f"WARN: image check failed for {path}: {e}")
        return path


def build(title, output, paragraphs, images, subtitle=None, captions=None, signature="火星前哨站"):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    head = doc.add_heading(level=1)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = head.add_run(title)
    set_font(run, "SimSun", "宋体", size=Pt(22), bold=True)
    run.font.color.rgb = TITLE_COLOR
    head.paragraph_format.space_after = Pt(6)

    if subtitle:
        for part in subtitle.split("\n"):
            part = part.strip()
            if not part:
                continue
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sub.add_run(part)
            set_font(sr, "SimSun", "宋体", size=Pt(12), bold=False)
            sr.font.color.rgb = SUBTITLE_COLOR

    captions = captions or []
    if len(paragraphs) != 7:
        print(f"WARN: expected 7 body paragraphs, got {len(paragraphs)}", file=sys.stderr)
    if len(images) != 3:
        print(f"WARN: expected 3 images, got {len(images)}", file=sys.stderr)
    if len(captions) != 3:
        print(f"WARN: expected 3 captions, got {len(captions)}", file=sys.stderr)
    for i, para in enumerate(paragraphs):
        if not para or "placeholder" in para.lower():
            print(f"WARN: paragraph[{i}] is empty or contains placeholder: {para!r}", file=sys.stderr)
    IMAGE_AFTER = {1, 3, 5}
    for idx, para in enumerate(paragraphs, start=1):
        p = doc.add_paragraph(para)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = LINE_SPACING
        for r in p.runs:
            set_font(r, "SimSun", "宋体", size=BODY_FONT_SIZE)
        if idx in IMAGE_AFTER:
            img_idx = (idx - 1) // 2
            if img_idx < len(images) and images[img_idx] and os.path.exists(images[img_idx]):
                img = ensure_docx_supported(images[img_idx])
                pic = doc.add_picture(img, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.paragraphs[-1].paragraph_format.space_before = Pt(6)
                doc.paragraphs[-1].paragraph_format.space_after = Pt(4)
                if img_idx < len(captions) and captions[img_idx]:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cap.add_run(captions[img_idx])
                    set_font(cr, "SimSun", "宋体", size=Pt(9), bold=False)
                    cr.font.color.rgb = CAPTION_COLOR

    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.paragraph_format.space_before = Pt(10)
    sr = sign.add_run(signature)
    set_font(sr, "SimSun", "宋体", size=Pt(12))
    sr.font.color.rgb = SUBTITLE_COLOR

    os.makedirs(os.path.dirname(os.path.abspath(output)) or ".", exist_ok=True)
    doc.save(output)
    print(f"OK docx -> {output}")


def main() -> int:
    p = argparse.ArgumentParser(description="Build a news article docx")
    p.add_argument("--title", required=True)
    p.add_argument("--output", required=True)
    p.add_argument("--body", required=True, help="JSON file: {\"subtitle\": ..., \"paragraphs\": [...], \"captions\": [...]}")
    p.add_argument("--images", nargs="*", default=[], help="local image paths (1-3)")
    p.add_argument("--signature", default="火星前哨站")
    args = p.parse_args()

    with open(args.body, "r", encoding="utf-8-sig") as f:
        data = json.load(f)
    build(
        args.title,
        args.output,
        data.get("paragraphs", []),
        args.images,
        subtitle=data.get("subtitle"),
        captions=data.get("captions", []),
        signature=args.signature,
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
