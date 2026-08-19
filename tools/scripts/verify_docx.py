import sys
import os
import re
import hashlib
import zipfile
import io
from datetime import datetime, timedelta
from difflib import SequenceMatcher
from docx import Document

EXPECTED_IMAGE_COUNT = 3
EXPECTED_BODY_PARAS = 7
MIN_BODY_CHARS = 1500
MAX_BODY_CHARS = 1900
BAD_MARKERS = ("captions_placeholder", "placeholder", "Lorem ipsum", "TBD", "待补充")
FORBIDDEN_TERMS = ("习近平", "习主席", "习总书记", "国家主席")
DHASH_DUP_THRESHOLD = 4

CROSS_DAY_WINDOW_DAYS = 7
TITLE_RATIO_DUP = 0.38
TITLE_LCS_DUP = 6
BODY_LCS_DUP = 45
BODY_RATIO_DUP = 0.30


def _dhash(data):
    from PIL import Image
    im = Image.open(io.BytesIO(data)).convert("L").resize((9, 8))
    px = list(im.getdata())
    bits = []
    for y in range(8):
        for x in range(8):
            bits.append("1" if px[y * 9 + x] > px[y * 9 + x + 1] else "0")
    return "".join(bits)


def _hamming(a, b):
    return sum(x != y for x, y in zip(a, b))


def _lcs(a, b):
    return SequenceMatcher(None, a, b, autojunk=False).find_longest_match(0, len(a), 0, len(b)).size


def _title_entities(title):
    return set(re.findall(r"[《〈【]([^》〉】]{2,20})[》〉】]", title))


def _recent_articles(base_dir, window_days, ref_path):
    ref_day = os.path.basename(os.path.dirname(os.path.abspath(ref_path)))
    if not re.match(r"^\d{4}-\d{2}-\d{2}$", ref_day):
        return []
    cutoff = (datetime.strptime(ref_day, "%Y-%m-%d") - timedelta(days=window_days)).strftime("%Y-%m-%d")
    out = []
    if not os.path.isdir(base_dir):
        return out
    for day in sorted(os.listdir(base_dir)):
        if not re.match(r"^\d{4}-\d{2}-\d{2}$", day) or day < cutoff or day > ref_day:
            continue
        for fn in sorted(os.listdir(os.path.join(base_dir, day))):
            if fn.endswith(".docx"):
                out.append(os.path.join(base_dir, day, fn))
    return out


def _common_blocks(a, b, min_size=10):
    sm = SequenceMatcher(None, a, b, autojunk=False)
    return [(blk.size, a[blk.a:blk.a + blk.size]) for blk in sm.get_matching_blocks() if blk.size >= min_size]


def check_cross_day_duplicates(path, window_days=CROSS_DAY_WINDOW_DAYS, base_dir=None):
    problems = []
    if window_days <= 0:
        return problems
    if base_dir is None:
        root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        base_dir = os.path.join(root, "article")
    try:
        doc = Document(path)
        paras = [p.text.strip() for p in doc.paragraphs]
    except Exception:
        return problems
    title = paras[0] if paras else ""
    body = "".join(p for p in paras[3:17] if p and not p.startswith("图｜") and p != "火星前哨站")
    entities = _title_entities(title)
    recents = [o for o in _recent_articles(base_dir, window_days, path) if os.path.abspath(o) != os.path.abspath(path)]
    pairs = []
    for other in recents:
        try:
            odoc = Document(other)
            oparas = [p.text.strip() for p in odoc.paragraphs]
            otitle = oparas[0] if oparas else ""
            obody = "".join(p for p in oparas[3:17] if p and not p.startswith("图｜") and p != "火星前哨站")
        except Exception:
            continue
        if not title or not otitle:
            continue
        pairs.append((other, otitle, obody))
    from collections import Counter
    block_count = Counter()
    for other, otitle, obody in pairs:
        for size, txt in _common_blocks(body, obody):
            block_count[txt] += 1
    for other, otitle, obody in pairs:
        tr = SequenceMatcher(None, title, otitle).ratio()
        tlcs = _lcs(title, otitle)
        blcs = max((size for size, txt in _common_blocks(body, obody) if block_count[txt] < 3), default=0)
        br = SequenceMatcher(None, body, obody).ratio()
        dup = None
        if entities and entities & _title_entities(otitle):
            dup = "same subject entity %s" % (entities & _title_entities(otitle))
        elif tr >= TITLE_RATIO_DUP and tlcs >= TITLE_LCS_DUP:
            dup = "title ratio %.2f / lcs %d" % (tr, tlcs)
        elif blcs >= BODY_LCS_DUP:
            dup = "body common substring %d chars" % blcs
        elif br >= BODY_RATIO_DUP:
            dup = "body ratio %.2f" % br
        if dup:
            problems.append("cross-day duplicate (within %dd, %s): %s" % (window_days, dup, os.path.basename(other)))
    return problems


def check_duplicate_images(path):
    problems = []
    items = []
    with zipfile.ZipFile(path) as z:
        for n in sorted(z.namelist()):
            if n.startswith("word/media/"):
                data = z.read(n)
                items.append((n, hashlib.md5(data).hexdigest(), data))
    for i in range(len(items)):
        for j in range(i + 1, len(items)):
            n1, h1, d1 = items[i]
            n2, h2, d2 = items[j]
            if h1 == h2:
                problems.append(f"duplicate image (identical): {n1} == {n2}")
                continue
            try:
                if _hamming(_dhash(d1), _dhash(d2)) <= DHASH_DUP_THRESHOLD:
                    problems.append(f"duplicate image (visually same, different size/encoding): {n1} ~ {n2}")
            except Exception:
                pass
    return problems


def verify(path, window_days=CROSS_DAY_WINDOW_DAYS):
    problems = []
    try:
        doc = Document(path)
    except Exception as e:
        return [f"cannot open docx: {e}"]

    paras = [p.text.strip() for p in doc.paragraphs]

    if len(paras) != 17:
        problems.append(f"paragraph count {len(paras)} != 17 (title+2subtitle+7body+3img+3caption+signature)")

    title = paras[0] if paras else ""
    if not title:
        problems.append("missing title (para 0 empty)")

    subtitle = paras[1] if len(paras) > 1 else ""
    if not subtitle or "国际观察" not in subtitle:
        problems.append(f"missing subtitle bar (para 1): {subtitle!r}")
    if len(paras) > 2 and not paras[2]:
        problems.append("missing subtitle summary line (para 2 empty)")

    body = paras[3:17]
    body = [t for t in body if t and "图｜" not in t and t not in ("火星前哨站",)]
    if len(body) != EXPECTED_BODY_PARAS:
        problems.append(f"body paragraphs {len(body)} != {EXPECTED_BODY_PARAS}")

    body_chars = sum(len(t) for t in body)
    if body_chars < MIN_BODY_CHARS or body_chars > MAX_BODY_CHARS:
        problems.append(f"body total chars {body_chars} outside range {MIN_BODY_CHARS}-{MAX_BODY_CHARS}")

    captions = [t for t in paras if t.startswith("图｜")]
    if len(captions) != 3:
        problems.append(f"captions {len(captions)} != 3")

    has_signature = any(t == "火星前哨站" for t in paras)
    if not has_signature:
        problems.append("missing signature '火星前哨站'")

    for marker in BAD_MARKERS:
        if any(marker in t for t in paras):
            problems.append(f"bad marker '{marker}' found in content")

    for term in FORBIDDEN_TERMS:
        if any(term in t for t in paras):
            problems.append(f"forbidden term '{term}' found in content (sensitive person)")

    n_images = len(doc.inline_shapes)
    if n_images != EXPECTED_IMAGE_COUNT:
        problems.append(f"inline images {n_images} != {EXPECTED_IMAGE_COUNT}")

    problems.extend(check_duplicate_images(path))
    problems.extend(check_cross_day_duplicates(path, window_days=window_days))

    return problems


def main() -> int:
    import argparse
    p = argparse.ArgumentParser(description="Verify news article docx")
    p.add_argument("docx", nargs="+", help="docx files to verify")
    p.add_argument("--recent-days", type=int, default=CROSS_DAY_WINDOW_DAYS,
                   help="cross-day duplicate window in days (0 to disable)")
    args = p.parse_args()
    rc = 0
    for path in args.docx:
        problems = verify(path, window_days=args.recent_days)
        if problems:
            rc = 1
            print(f"FAIL {path}:")
            for prob in problems:
                print(f"  - {prob}")
        else:
            print(f"PASS {path}")
    return rc


if __name__ == "__main__":
    sys.exit(main())
